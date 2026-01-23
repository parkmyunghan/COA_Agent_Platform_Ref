# ui/components/report_engine.py
# -*- coding: utf-8 -*-
"""
보고서 생성 엔진
다양한 형식(PDF, Word) 및 타입 지원
"""
import streamlit as st
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, List
import json


class ReportEngine:
    """보고서 생성 엔진"""
    
    def __init__(self, config: Optional[Dict] = None):
        self.config = config or {}
        self.reports_dir = Path("./reports/generated")
        self.reports_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_report(
        self,
        report_type: str,  # "situation", "coa", "rationale", "execution"
        data: Dict,
        format: str = "pdf",  # "pdf", "docx", "html", "xlsx"
        include_charts: bool = True,
        include_details: bool = True,
        include_appendix: bool = False
    ) -> Optional[str]:
        """보고서 생성"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if format == "pdf":
            return self._generate_pdf(report_type, data, timestamp, include_charts, include_details, include_appendix)
        elif format == "docx":
            return self._generate_docx(report_type, data, timestamp, include_charts, include_details, include_appendix)
        elif format == "html":
            return self._generate_html(report_type, data, timestamp, include_charts, include_details, include_appendix)
        elif format == "xlsx":
            return self._generate_excel(report_type, data, timestamp, include_charts, include_details, include_appendix)
        else:
            st.error(f"지원하지 않는 형식: {format}")
            return None
    
    def generate_situation_report(self, situation_info: Dict, format: str = "pdf") -> Optional[str]:
        """상황 분석 보고서 생성"""
        data = {
            "situation_info": situation_info,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "report_type": "situation"
        }
        return self.generate_report("situation", data, format)
    
    def generate_coa_report(self, agent_result: Dict, format: str = "pdf") -> Optional[str]:
        """방책 추천 보고서 생성"""
        recommendations = agent_result.get("recommendations", [])
        situation_info = agent_result.get("situation_info", {})
        situation_analysis = agent_result.get("situation_analysis", {})
        reasoning_process = agent_result.get("reasoning_process", {})
        
        # 자연어 설명 가져오기 (세션 상태에서)
        try:
            import streamlit as st
            # 여러 가능한 키에서 자연어 설명 찾기
            natural_language_explanation = (
                st.session_state.get("reasoning_nl_explanation", "") or
                st.session_state.get("nl_explanation", "") or
                agent_result.get("natural_language_explanation", "")
            )
        except:
            # Streamlit 컨텍스트가 없는 경우 (예: 테스트 환경)
            natural_language_explanation = agent_result.get("natural_language_explanation", "")
        
        data = {
            "recommendations": recommendations[:3],  # 상위 3개
            "reasoning_process": reasoning_process,
            "situation_info": situation_info,
            "situation_analysis": situation_analysis,  # 상황 분석 정보 추가
            "natural_language_explanation": natural_language_explanation,  # 자연어 설명 추가
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "report_type": "coa"
        }
        return self.generate_report("coa", data, format)
    
    def generate_rationale_report(self, agent_result: Dict, format: str = "pdf") -> Optional[str]:
        """의사결정 근거 보고서 생성"""
        data = {
            "situation_info": agent_result.get("situation_info", {}),
            "reasoning_steps": self._extract_reasoning_steps(agent_result),
            "score_calculation": self._extract_score_calculation(agent_result),
            "recommendations": agent_result.get("recommendations", []),
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "report_type": "rationale"
        }
        return self.generate_report("rationale", data, format)
    
    def generate_execution_plan(self, recommendation: Dict, situation_info: Dict = None, format: str = "pdf") -> Optional[str]:
        """실행 계획서 생성"""
        data = {
            "coa": recommendation,
            "execution_steps": self._generate_execution_steps(recommendation),
            "resource_requirements": self._extract_resource_requirements(recommendation),
            "risk_assessment": self._assess_risks(recommendation),
            "situation_info": situation_info or {},
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "report_type": "execution"
        }
        return self.generate_report("execution", data, format)
    
    def _register_korean_font(self):
        """한글 폰트 등록 (Windows 시스템)"""
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import platform
            
            # Windows 시스템 폰트 경로
            if platform.system() == 'Windows':
                font_paths = [
                    'C:/Windows/Fonts/malgun.ttf',      # 맑은 고딕
                    'C:/Windows/Fonts/malgunbd.ttf',    # 맑은 고딕 Bold
                    'C:/Windows/Fonts/gulim.ttc',       # 굴림 (TTC)
                    'C:/Windows/Fonts/batang.ttc',      # 바탕 (TTC)
                ]
                
                # TTF 파일 찾기 및 등록
                for font_path in font_paths:
                    if font_path.endswith('.ttf') and Path(font_path).exists():
                        try:
                            pdfmetrics.registerFont(TTFont('KoreanFont', font_path))
                            # Bold 폰트도 등록 시도
                            if 'bd' in font_path.lower() or 'bold' in font_path.lower():
                                pdfmetrics.registerFont(TTFont('KoreanFontBold', font_path))
                            else:
                                # 일반 폰트를 Bold로도 사용
                                pdfmetrics.registerFont(TTFont('KoreanFontBold', font_path))
                            return 'KoreanFont'
                        except Exception:
                            continue
                
                # TTC 파일 처리 (fonttools 필요)
                try:
                    from fontTools.ttLib import TTFont as FontToolsTTFont
                    for font_path in font_paths:
                        if font_path.endswith('.ttc') and Path(font_path).exists():
                            try:
                                # TTC에서 첫 번째 폰트 추출
                                ttc = FontToolsTTFont(font_path, fontNumber=0)
                                # 임시 TTF 파일로 저장
                                temp_dir = Path('./temp_fonts')
                                temp_dir.mkdir(exist_ok=True)
                                temp_ttf = temp_dir / 'korean_font_temp.ttf'
                                ttc.save(str(temp_ttf))
                                pdfmetrics.registerFont(TTFont('KoreanFont', str(temp_ttf)))
                                pdfmetrics.registerFont(TTFont('KoreanFontBold', str(temp_ttf)))
                                return 'KoreanFont'
                            except Exception:
                                continue
                except ImportError:
                    # fonttools가 없으면 TTC 파일은 건너뜀
                    pass
            
            # 폰트를 찾지 못한 경우 경고
            import streamlit as st
            st.warning("한글 폰트를 찾을 수 없습니다. 기본 폰트를 사용합니다. (한글이 제대로 표시되지 않을 수 있습니다)")
            return 'Helvetica'
        except Exception as e:
            import streamlit as st
            st.warning(f"폰트 등록 실패: {e}. 기본 폰트를 사용합니다.")
            return 'Helvetica'
    
    def _generate_pdf(
        self,
        report_type: str,
        data: Dict,
        timestamp: str,
        include_charts: bool,
        include_details: bool,
        include_appendix: bool
    ) -> Optional[str]:
        """PDF 생성"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
            
            # 한글 폰트 등록
            korean_font = self._register_korean_font()
            korean_font_bold = 'KoreanFontBold' if korean_font == 'KoreanFont' else 'Helvetica-Bold'
            
            filename = f"{report_type}_report_{timestamp}.pdf"
            output_path = self.reports_dir / filename
            
            doc = SimpleDocTemplate(str(output_path), pagesize=A4)
            story = []
            
            # 스타일 설정 (한글 폰트 사용)
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.HexColor('#1f77b4'),
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName=korean_font  # 한글 폰트 추가
            )
            
            # 한글 폰트를 사용하는 Normal 스타일
            normal_style = ParagraphStyle(
                'KoreanNormal',
                parent=styles['Normal'],
                fontName=korean_font,
                fontSize=10
            )
            
            # 한글 폰트를 사용하는 Heading 스타일들
            heading1_style = ParagraphStyle(
                'KoreanHeading1',
                parent=styles['Heading1'],
                fontName=korean_font,
                fontSize=18
            )
            
            heading2_style = ParagraphStyle(
                'KoreanHeading2',
                parent=styles['Heading2'],
                fontName=korean_font,
                fontSize=14
            )
            
            heading3_style = ParagraphStyle(
                'KoreanHeading3',
                parent=styles['Heading3'],
                fontName=korean_font,
                fontSize=12
            )
            
            # 표지
            story.append(Paragraph(self._get_report_title(report_type), title_style))
            story.append(Spacer(1, 0.3*inch))
            story.append(Paragraph(f"<b>생성일시:</b> {data.get('timestamp', '')}", normal_style))
            story.append(PageBreak())
            
            # 한글 폰트 스타일을 딕셔너리로 전달
            korean_styles = {
                'Normal': normal_style,
                'Heading1': heading1_style,
                'Heading2': heading2_style,
                'Heading3': heading3_style,
                'KoreanFont': korean_font,
                'KoreanFontBold': korean_font_bold
            }
            
            # 보고서 타입별 내용 생성
            if report_type == "coa":
                story.extend(self._build_coa_pdf_content(data, korean_styles, include_charts, include_details, include_appendix))
            elif report_type == "situation":
                story.extend(self._build_situation_pdf_content(data, korean_styles, include_details))
            elif report_type == "rationale":
                story.extend(self._build_rationale_pdf_content(data, korean_styles, include_details))
            elif report_type == "execution":
                story.extend(self._build_execution_pdf_content(data, korean_styles, include_details))
            
            # PDF 빌드
            doc.build(story)
            return str(output_path)
            
        except ImportError:
            st.error("ReportLab이 설치되지 않았습니다. pip install reportlab")
            return None
        except Exception as e:
            st.error(f"PDF 생성 실패: {e}")
            import traceback
            st.code(traceback.format_exc())
            return None
    
    def _generate_docx(
        self,
        report_type: str,
        data: Dict,
        timestamp: str,
        include_charts: bool,
        include_details: bool,
        include_appendix: bool
    ) -> Optional[str]:
        """Word 문서 생성"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            filename = f"{report_type}_report_{timestamp}.docx"
            output_path = self.reports_dir / filename
            
            doc = Document()
            
            # 제목
            title = doc.add_heading(self._get_report_title(report_type), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 생성일시
            doc.add_paragraph(f"생성일시: {data.get('timestamp', '')}")
            doc.add_page_break()
            
            # 보고서 타입별 내용 생성
            if report_type == "coa":
                self._build_coa_docx_content(doc, data, include_charts, include_details, include_appendix)
            elif report_type == "situation":
                self._build_situation_docx_content(doc, data, include_details)
            elif report_type == "rationale":
                self._build_rationale_docx_content(doc, data, include_details)
            elif report_type == "execution":
                self._build_execution_docx_content(doc, data, include_details)
            
            doc.save(str(output_path))
            return str(output_path)
            
        except ImportError:
            st.error("python-docx가 설치되지 않았습니다. pip install python-docx")
            return None
        except Exception as e:
            st.error(f"Word 문서 생성 실패: {e}")
            import traceback
            st.code(traceback.format_exc())
            return None
    
    def _generate_html(
        self,
        report_type: str,
        data: Dict,
        timestamp: str,
        include_charts: bool,
        include_details: bool,
        include_appendix: bool
    ) -> Optional[str]:
        """HTML 생성 (Jinja2 템플릿 사용)"""
        try:
            from jinja2 import Environment, FileSystemLoader, Template
            
            # 템플릿 디렉토리 설정
            template_dir = Path(__file__).parent.parent.parent / "reports" / "templates"
            style_dir = Path(__file__).parent.parent.parent / "reports" / "styles"
            
            # Jinja2 환경 설정
            env = Environment(
                loader=FileSystemLoader(str(template_dir)),
                autoescape=True
            )
            
            # CSS 파일 읽기
            css_file = style_dir / "military_report.css"
            css_style = ""
            if css_file.exists():
                with open(css_file, 'r', encoding='utf-8') as f:
                    css_style = f.read()
            
            # 템플릿 파일 선택
            template_mapping = {
                "situation": "situation_analysis_template.html",
                "coa": "coa_recommendation_template.html",
                "rationale": "decision_rationale_template.html",
                "execution": "execution_plan_template.html"
            }
            
            template_file = template_mapping.get(report_type, "coa_recommendation_template.html")
            
            # 템플릿 로드 및 렌더링
            template = env.get_template(template_file)
            
            # 데이터 준비
            template_data = {
                "title": self._get_report_title(report_type),
                "timestamp": data.get('timestamp', timestamp),
                "css_style": css_style,
                "include_charts": include_charts,
                "include_details": include_details,
                "include_appendix": include_appendix,
                **data
            }
            
            html_content = template.render(**template_data)
            
            # HTML 파일 저장
            filename = f"{report_type}_report_{timestamp}.html"
            output_path = self.reports_dir / filename
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return str(output_path)
            
        except ImportError:
            st.error("Jinja2가 설치되지 않았습니다. `pip install jinja2`를 실행하세요.")
            return None
        except Exception as e:
            st.error(f"HTML 생성 실패: {e}")
            import traceback
            st.code(traceback.format_exc())
            return None
    
    def _get_report_title(self, report_type: str) -> str:
        """보고서 제목"""
        titles = {
            "situation": "상황 분석 보고서",
            "coa": "방책 추천 보고서",
            "rationale": "의사결정 근거 보고서",
            "execution": "실행 계획서"
        }
        return titles.get(report_type, "보고서")
    
    def _build_coa_pdf_content(self, data: Dict, styles, include_charts: bool, include_details: bool, include_appendix: bool):
        """방책 추천 보고서 PDF 내용"""
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        
        story = []
        recommendations = data.get("recommendations", [])
        situation_info = data.get("situation_info", {})
        situation_analysis = data.get("situation_analysis", {})
        
        # 한글 폰트 정보 추출
        korean_font = styles.get('KoreanFont', 'Helvetica')
        korean_font_bold = styles.get('KoreanFontBold', 'Helvetica-Bold')
        normal_style = styles.get('Normal')
        heading1_style = styles.get('Heading1')
        heading2_style = styles.get('Heading2')
        heading3_style = styles.get('Heading3')
        
        # 위협 상황 상세 설명 섹션 추가
        story.append(Paragraph("<b>위협 상황 상세</b>", heading2_style))
        if situation_info:
            # 기본 정보 표
            situation_data = []
            
            # 주요 필드들 추가
            threat_type = situation_info.get('위협유형', situation_info.get('위협유형', 'N/A'))
            if threat_type and threat_type != 'N/A':
                situation_data.append([
                    Paragraph("<b>위협 유형</b>", normal_style),
                    Paragraph(str(threat_type), normal_style)
                ])
            
            threat_level = situation_info.get('심각도', situation_info.get('위협수준', 'N/A'))
            if threat_level and threat_level != 'N/A':
                if isinstance(threat_level, (int, float)):
                    threat_level_text = f"{threat_level:.1f}%"
                else:
                    threat_level_text = str(threat_level)
                situation_data.append([
                    Paragraph("<b>위협 수준</b>", normal_style),
                    Paragraph(threat_level_text, normal_style)
                ])
            
            location = situation_info.get('발생장소', situation_info.get('장소', 'N/A'))
            if location and location != 'N/A':
                situation_data.append([
                    Paragraph("<b>발생 장소</b>", normal_style),
                    Paragraph(str(location), normal_style)
                ])
            
            detection_time = situation_info.get('탐지시각', situation_info.get('탐지시간', ''))
            if detection_time:
                situation_data.append([
                    Paragraph("<b>탐지 시각</b>", normal_style),
                    Paragraph(str(detection_time), normal_style)
                ])
            
            evidence = situation_info.get('근거', situation_info.get('증거', ''))
            if evidence:
                situation_data.append([
                    Paragraph("<b>탐지 근거</b>", normal_style),
                    Paragraph(str(evidence), normal_style)
                ])
            
            threat_id = situation_info.get('위협ID', situation_info.get('ID', ''))
            if threat_id:
                situation_data.append([
                    Paragraph("<b>위협 ID</b>", normal_style),
                    Paragraph(str(threat_id), normal_style)
                ])
            
            # 추가 컨텍스트 정보
            additional_context = situation_info.get('additional_context', '')
            if additional_context:
                situation_data.append([
                    Paragraph("<b>추가 정보</b>", normal_style),
                    Paragraph(str(additional_context), normal_style)
                ])
            
            if situation_data:
                situation_table = Table(situation_data, colWidths=[2*inch, 4*inch])
                situation_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, 0), korean_font_bold),
                    ('FONTNAME', (0, 1), (-1, -1), korean_font),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                    ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey)
                ]))
                story.append(situation_table)
                story.append(Spacer(1, 0.2*inch))
            
            # 상황 분석 결과가 있으면 추가 설명
            if situation_analysis:
                # RAG 검색 결과 요약
                rag_results = situation_analysis.get("rag_results", [])
                if rag_results:
                    story.append(Paragraph("<b>관련 문서 정보</b>", heading3_style))
                    rag_summary = f"총 {len(rag_results)}개의 관련 문서가 검색되었습니다. "
                    if len(rag_results) > 0:
                        top_rag = rag_results[0]
                        rag_text = top_rag.get("text", "")[:200] if isinstance(top_rag.get("text"), str) else ""
                        if rag_text:
                            rag_summary += f"주요 관련 내용: {rag_text}..."
                    story.append(Paragraph(rag_summary, normal_style))
                    story.append(Spacer(1, 0.15*inch))
                
                # 관련 엔티티 정보
                related_entities = situation_analysis.get("related_entities", [])
                if related_entities:
                    story.append(Paragraph("<b>관련 엔티티</b>", heading3_style))
                    entities_text = f"총 {len(related_entities)}개의 관련 엔티티가 식별되었습니다: "
                    entity_names = [str(e.get("name", e.get("label", ""))) for e in related_entities[:5]]
                    entities_text += ", ".join([e for e in entity_names if e])
                    if len(related_entities) > 5:
                        entities_text += f" 외 {len(related_entities) - 5}개"
                    story.append(Paragraph(entities_text, normal_style))
                    story.append(Spacer(1, 0.15*inch))
                
                # 관계 체인 정보
                chain_info = situation_analysis.get("chain_info", {})
                if chain_info and chain_info.get("chains"):
                    chains = chain_info.get("chains", [])
                    if chains:
                        story.append(Paragraph("<b>관계 체인 분석</b>", heading3_style))
                        chain_summary = f"총 {len(chains)}개의 관계 체인이 발견되었습니다. "
                        if len(chains) > 0:
                            top_chain = chains[0]
                            chain_desc = top_chain.get("description", top_chain.get("summary", ""))
                            if chain_desc:
                                chain_summary += f"주요 관계: {chain_desc[:150]}..."
                        story.append(Paragraph(chain_summary, normal_style))
                        story.append(Spacer(1, 0.15*inch))
        
        story.append(Spacer(1, 0.3*inch))
        
        # Executive Summary
        story.append(Paragraph("<b>Executive Summary</b>", heading2_style))
        if situation_info:
            summary_text = f"위협 상황: {situation_info.get('위협유형', 'N/A')}<br/>"
            summary_text += f"위협 수준: {situation_info.get('심각도', 'N/A')}<br/>"
            summary_text += f"발생 장소: {situation_info.get('발생장소', 'N/A')}"
            story.append(Paragraph(summary_text, normal_style))
        story.append(Spacer(1, 0.3*inch))
        
        # 추천 방책 요약
        story.append(Paragraph("<b>추천 방책 요약</b>", heading2_style))
        if recommendations:
            # 헤더는 Paragraph로 생성
            rec_data = [[Paragraph("<b>순위</b>", normal_style), 
                         Paragraph("<b>방책명</b>", normal_style), 
                         Paragraph("<b>적합도 점수</b>", normal_style)]]
            for i, rec in enumerate(recommendations, 1):
                coa_name = rec.get('coa_name', f'방책 {i}')
                rec_data.append([
                    Paragraph(str(i), normal_style),
                    Paragraph(coa_name, normal_style),  # Paragraph 객체 사용하여 자동 줄바꿈
                    Paragraph(f"{rec.get('score', 0):.2f}", normal_style)
                ])
            
            rec_table = Table(rec_data, colWidths=[1*inch, 3*inch, 1.5*inch])
            rec_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # 상단 정렬 추가
                ('FONTNAME', (0, 0), (-1, 0), korean_font_bold),  # 한글 폰트 사용
                ('FONTNAME', (0, 1), (-1, -1), korean_font),  # 본문도 한글 폰트 사용
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 1), (-1, -1), 6),  # 상단 패딩 추가
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),  # 하단 패딩 추가
                ('LEFTPADDING', (0, 0), (-1, -1), 6),  # 좌측 패딩 추가
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),  # 우측 패딩 추가
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(rec_table)
            story.append(Spacer(1, 0.3*inch))
        
        # 자연어 설명 섹션 추가
        natural_language_explanation = data.get("natural_language_explanation", "")
        if natural_language_explanation and natural_language_explanation.strip():
            story.append(Paragraph("<b>방책 추천 설명</b>", heading2_style))
            # 마크다운 형식의 줄바꿈을 HTML <br/>로 변환
            explanation_text = natural_language_explanation.replace('\n\n', '<br/><br/>').replace('\n', '<br/>')
            # 마크다운 헤더를 HTML로 변환 (더 정확한 변환)
            import re
            # 헤더 변환 (#, ##, ###)
            explanation_text = re.sub(r'^### (.+?)$', r'<b>\1</b>', explanation_text, flags=re.MULTILINE)
            explanation_text = re.sub(r'^## (.+?)$', r'<b>\1</b>', explanation_text, flags=re.MULTILINE)
            explanation_text = re.sub(r'^# (.+?)$', r'<b>\1</b>', explanation_text, flags=re.MULTILINE)
            # 간단한 마크다운 볼드/이탤릭 처리
            explanation_text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', explanation_text)
            explanation_text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', explanation_text)
            
            story.append(Paragraph(explanation_text, normal_style))
            story.append(Spacer(1, 0.3*inch))
        
        # 상세 정보
        if include_details and recommendations:
            story.append(PageBreak())
            story.append(Paragraph("<b>추천 방책 상세</b>", heading1_style))
            
            for i, rec in enumerate(recommendations, 1):
                story.append(Paragraph(f"<b>{i}. {rec.get('coa_name', f'방책 {i}')}</b>", heading2_style))
                story.append(Paragraph(f"적합도 점수: {rec.get('score', 0):.2f}", normal_style))
                
                if rec.get('description'):
                    story.append(Paragraph(f"<b>설명:</b> {rec.get('description', '')}", normal_style))
                
                # 점수 breakdown
                score_breakdown = rec.get('score_breakdown', {})
                if score_breakdown:
                    story.append(Paragraph("<b>점수 Breakdown:</b>", heading3_style))
                    # Paragraph 객체로 변환하여 줄바꿈 지원
                    breakdown_data = [[Paragraph("<b>요소</b>", normal_style), 
                                        Paragraph("<b>점수</b>", normal_style)]]
                    for key, value in score_breakdown.items():
                        # METT-C 점수는 별도 섹션에서 처리
                        if key == 'mett_c':
                            continue
                        breakdown_data.append([
                            Paragraph(str(key), normal_style),
                            Paragraph(f"{value:.2f}", normal_style)
                        ])
                    
                    breakdown_table = Table(breakdown_data, colWidths=[2*inch, 1*inch])
                    breakdown_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # 상단 정렬 추가
                        ('FONTNAME', (0, 0), (-1, 0), korean_font_bold),  # 한글 폰트 사용
                        ('FONTNAME', (0, 1), (-1, -1), korean_font),  # 본문도 한글 폰트 사용
                        ('TOPPADDING', (0, 1), (-1, -1), 6),  # 상단 패딩 추가
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),  # 하단 패딩 추가
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),  # 좌측 패딩 추가
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),  # 우측 패딩 추가
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(breakdown_table)
                
                # METT-C 점수 (있는 경우)
                mett_c_scores = score_breakdown.get("mett_c") or rec.get("mett_c")
                if mett_c_scores:
                    story.append(Spacer(1, 0.15*inch))
                    story.append(Paragraph("<b>METT-C 종합 평가:</b>", heading3_style))
                    
                    mett_c_data = [[Paragraph("<b>요소</b>", normal_style), 
                                   Paragraph("<b>점수</b>", normal_style),
                                   Paragraph("<b>해석</b>", normal_style)]]
                    
                    mett_c_elements = {
                        "mission": ("🎯 임무", 0.8),
                        "enemy": ("⚠️ 적군", 0.6),
                        "terrain": ("🌍 지형", 0.6),
                        "troops": ("👥 부대", 0.6),
                        "civilian": ("🏘️ 민간인", 0.3),
                        "time": ("⏰ 시간", 0.5)
                    }
                    
                    for key, (label, threshold) in mett_c_elements.items():
                        score = mett_c_scores.get(key, 0)
                        if score >= threshold:
                            interpretation = "양호"
                        elif score >= threshold * 0.5:
                            interpretation = "보통"
                        else:
                            interpretation = "부족"
                        
                        # 민간인/시간 특별 표시
                        if key == "civilian" and score < 0.3:
                            interpretation = "⚠️ 민간인 보호 낮음"
                        elif key == "time" and score == 0.0:
                            interpretation = "❌ 시간 제약 위반"
                        elif key == "time" and score < 0.5:
                            interpretation = "⚠️ 시간 제약 주의"
                        
                        mett_c_data.append([
                            Paragraph(label, normal_style),
                            Paragraph(f"{score:.3f}", normal_style),
                            Paragraph(interpretation, normal_style)
                        ])
                    
                    mett_c_table = Table(mett_c_data, colWidths=[1.5*inch, 1*inch, 2*inch])
                    mett_c_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#00b894')),  # 녹색 배경
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('FONTNAME', (0, 0), (-1, 0), korean_font_bold),
                        ('FONTNAME', (0, 1), (-1, -1), korean_font),
                        ('FONTSIZE', (0, 0), (-1, 0), 11),
                        ('FONTSIZE', (0, 1), (-1, -1), 10),
                        ('TOPPADDING', (0, 1), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(mett_c_table)
                    
                    # METT-C 종합 점수
                    mett_c_total = mett_c_scores.get("total", 0)
                    story.append(Spacer(1, 0.1*inch))
                    story.append(Paragraph(f"<b>METT-C 종합 점수:</b> {mett_c_total:.3f}", normal_style))
                
                story.append(Spacer(1, 0.2*inch))
        
        return story
    
    def _build_coa_docx_content(self, doc, data: Dict, include_charts: bool, include_details: bool, include_appendix: bool):
        """방책 추천 보고서 Word 내용"""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        recommendations = data.get("recommendations", [])
        situation_info = data.get("situation_info", {})
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        if situation_info:
            doc.add_paragraph(f"위협 상황: {situation_info.get('위협유형', 'N/A')}")
            doc.add_paragraph(f"위협 수준: {situation_info.get('심각도', 'N/A')}")
            doc.add_paragraph(f"발생 장소: {situation_info.get('발생장소', 'N/A')}")
        
        # 추천 방책 요약
        doc.add_heading('추천 방책 요약', 1)
        if recommendations:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '순위'
            hdr_cells[1].text = '방책명'
            hdr_cells[2].text = '적합도 점수'
            
            for i, rec in enumerate(recommendations, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = rec.get('coa_name', f'방책 {i}')
                row_cells[2].text = f"{rec.get('score', 0):.2f}"
        
        # 자연어 설명 섹션 추가
        natural_language_explanation = data.get("natural_language_explanation", "")
        if natural_language_explanation and natural_language_explanation.strip():
            doc.add_heading('방책 추천 설명', 1)
            # 마크다운 형식의 텍스트를 단락으로 분리하여 추가
            paragraphs = natural_language_explanation.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # 마크다운 헤더 처리
                    if para.strip().startswith('###'):
                        heading_text = para.strip().replace('###', '').strip()
                        doc.add_heading(heading_text, 3)
                    elif para.strip().startswith('##'):
                        heading_text = para.strip().replace('##', '').strip()
                        doc.add_heading(heading_text, 2)
                    elif para.strip().startswith('#'):
                        heading_text = para.strip().replace('#', '').strip()
                        doc.add_heading(heading_text, 1)
                    else:
                        # 일반 텍스트는 볼드 처리된 부분을 유지하면서 추가
                        doc.add_paragraph(para.strip())
            doc.add_paragraph()
        
        # 상세 정보
        if include_details and recommendations:
            doc.add_page_break()
            doc.add_heading('추천 방책 상세', 1)
            
            for i, rec in enumerate(recommendations, 1):
                doc.add_heading(f"{i}. {rec.get('coa_name', f'방책 {i}')}", 2)
                doc.add_paragraph(f"적합도 점수: {rec.get('score', 0):.2f}")
                
                if rec.get('description'):
                    doc.add_paragraph(f"설명: {rec.get('description', '')}")
                
                # 점수 breakdown
                if rec.get('score_breakdown'):
                    doc.add_heading('점수 Breakdown', 3)
                    breakdown_table = doc.add_table(rows=1, cols=2)
                    breakdown_table.style = 'Light List Accent 1'
                    hdr_cells = breakdown_table.rows[0].cells
                    hdr_cells[0].text = '요소'
                    hdr_cells[1].text = '점수'
                    
                    for key, value in rec.get('score_breakdown', {}).items():
                        row_cells = breakdown_table.add_row().cells
                        row_cells[0].text = key
                        row_cells[1].text = f"{value:.2f}"
    
    def _build_situation_pdf_content(self, data: Dict, styles, include_details: bool):
        """상황 분석 보고서 PDF 내용"""
        from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        story = []
        
        # 한글 폰트 정보 추출
        korean_font = styles.get('KoreanFont', 'Helvetica')
        korean_font_bold = styles.get('KoreanFontBold', 'Helvetica-Bold')
        normal_style = styles.get('Normal')
        heading2_style = styles.get('Heading2')
        
        # 상황 정보 표시
        situation_info = data.get("situation_info", {})
        if situation_info:
            story.append(Paragraph("<b>상황 정보</b>", heading2_style))
            situation_data = []
            for key, value in situation_info.items():
                situation_data.append([str(key), str(value)])
            
            if situation_data:
                situation_table = Table(situation_data, colWidths=[2*inch, 4*inch])
                situation_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), korean_font_bold),
                    ('FONTNAME', (0, 1), (-1, -1), korean_font),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(situation_table)
                story.append(Spacer(1, 0.3*inch))
        
        return story
    
    def _build_situation_docx_content(self, doc, data: Dict, include_details: bool):
        """상황 분석 보고서 Word 내용"""
        situation_info = data.get("situation_info", {})
        doc.add_heading('위협 상황 상세', 1)
        for key, value in situation_info.items():
            doc.add_paragraph(f"{key}: {value}")
    
    def _build_rationale_pdf_content(self, data: Dict, styles, include_details: bool):
        """의사결정 근거 보고서 PDF 내용"""
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.units import inch
        
        story = []
        
        # 한글 폰트 정보 추출
        korean_font = styles.get('KoreanFont', 'Helvetica')
        normal_style = styles.get('Normal')
        heading2_style = styles.get('Heading2')
        
        story.append(Paragraph("<b>의사결정 근거</b>", heading2_style))
        story.append(Paragraph("상세 내용은 추후 구현됩니다.", normal_style))
        story.append(Spacer(1, 0.3*inch))
        
        return story
    
    def _build_rationale_docx_content(self, doc, data: Dict, include_details: bool):
        """의사결정 근거 보고서 Word 내용"""
        doc.add_heading('의사결정 근거', 1)
        doc.add_paragraph("상세 내용은 추후 구현됩니다.")
    
    def _build_execution_pdf_content(self, data: Dict, styles, include_details: bool):
        """실행 계획서 PDF 내용"""
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.units import inch
        
        story = []
        
        # 한글 폰트 정보 추출
        korean_font = styles.get('KoreanFont', 'Helvetica')
        normal_style = styles.get('Normal')
        heading2_style = styles.get('Heading2')
        heading3_style = styles.get('Heading3')
        
        story.append(Paragraph("<b>실행 계획</b>", heading2_style))
        coa = data.get("coa", {})
        story.append(Paragraph(f"<b>방책명:</b> {coa.get('coa_name', 'N/A')}", normal_style))
        
        execution_steps = data.get("execution_steps", [])
        if execution_steps:
            story.append(Paragraph("<b>단계별 실행 계획:</b>", heading3_style))
            for i, step in enumerate(execution_steps, 1):
                story.append(Paragraph(f"{i}. {step}", normal_style))
        
        story.append(Spacer(1, 0.3*inch))
        
        return story
    
    def _build_execution_docx_content(self, doc, data: Dict, include_details: bool):
        """실행 계획서 Word 내용"""
        coa = data.get("coa", {})
        doc.add_heading('실행 계획', 1)
        doc.add_paragraph(f"방책명: {coa.get('coa_name', 'N/A')}")
        
        execution_steps = data.get("execution_steps", [])
        if execution_steps:
            doc.add_heading('단계별 실행 계획', 2)
            for i, step in enumerate(execution_steps, 1):
                doc.add_paragraph(f"{i}. {step}", style='List Number')
    
    # Helper methods
    def _extract_reasoning_steps(self, agent_result: Dict) -> Dict:
        """추론 과정 추출"""
        return agent_result.get("reasoning_process", {})
    
    def _extract_score_calculation(self, agent_result: Dict) -> Dict:
        """점수 계산 상세 추출"""
        recommendations = agent_result.get("recommendations", [])
        if recommendations:
            return recommendations[0].get("score_breakdown", {})
        return {}
    
    def _generate_execution_steps(self, recommendation: Dict) -> List[str]:
        """실행 단계 생성"""
        return [
            "Phase 1: 초기 배치 및 상황 파악",
            "Phase 2: 자원 배치 및 통신망 구축",
            "Phase 3: 방책 실행 및 모니터링",
            "Phase 4: 결과 평가 및 조정"
        ]
    
    def _extract_resource_requirements(self, recommendation: Dict) -> Dict:
        """필요 자원 추출"""
        return recommendation.get("resource_requirements", {})
    
    def _assess_risks(self, recommendation: Dict) -> Dict:
        """위험 평가"""
        return recommendation.get("risk_assessment", {})
    
    def _generate_excel(
        self,
        report_type: str,
        data: Dict,
        timestamp: str,
        include_charts: bool,
        include_details: bool,
        include_appendix: bool
    ) -> Optional[str]:
        """Excel 생성 (openpyxl 사용)"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            
            filename = f"{report_type}_report_{timestamp}.xlsx"
            output_path = self.reports_dir / filename
            
            wb = Workbook()
            ws = wb.active
            ws.title = "보고서 요약"
            
            # 스타일 정의
            header_fill = PatternFill(start_color="1f77b4", end_color="1f77b4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=12)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            center_align = Alignment(horizontal='center', vertical='center')
            
            # 제목
            ws['A1'] = self._get_report_title(report_type)
            ws['A1'].font = Font(bold=True, size=16)
            ws.merge_cells('A1:D1')
            ws['A1'].alignment = center_align
            
            row = 3
            
            # 생성일시
            ws[f'A{row}'] = "생성일시:"
            ws[f'B{row}'] = data.get('timestamp', timestamp)
            row += 2
            
            # 보고서 타입별 내용 생성
            if report_type == "coa":
                row = self._build_coa_excel_content(ws, data, row, header_fill, header_font, border, center_align, include_details)
            elif report_type == "situation":
                row = self._build_situation_excel_content(ws, data, row, header_fill, header_font, border, center_align)
            elif report_type == "rationale":
                row = self._build_rationale_excel_content(ws, data, row, header_fill, header_font, border, center_align)
            elif report_type == "execution":
                row = self._build_execution_excel_content(ws, data, row, header_fill, header_font, border, center_align)
            
            # 열 너비 자동 조정
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if cell.value and len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # 파일 저장
            wb.save(str(output_path))
            return str(output_path)
            
        except ImportError:
            st.error("openpyxl이 설치되지 않았습니다. `pip install openpyxl`을 실행하세요.")
            return None
        except Exception as e:
            st.error(f"Excel 생성 실패: {e}")
            import traceback
            st.code(traceback.format_exc())
            return None
    
    def _build_coa_excel_content(self, ws, data: Dict, start_row: int, header_fill, header_font, border, center_align, include_details: bool) -> int:
        """방책 추천 보고서 Excel 내용"""
        row = start_row
        
        # Executive Summary
        ws[f'A{row}'] = "Executive Summary"
        ws[f'A{row}'].font = Font(bold=True, size=14)
        row += 1
        
        situation_info = data.get("situation_info", {})
        if situation_info:
            ws[f'A{row}'] = "위협 상황:"
            ws[f'B{row}'] = situation_info.get('위협유형', 'N/A')
            row += 1
            ws[f'A{row}'] = "위협 수준:"
            ws[f'B{row}'] = situation_info.get('심각도', 'N/A')
            row += 1
            ws[f'A{row}'] = "발생 장소:"
            ws[f'B{row}'] = situation_info.get('발생장소', 'N/A')
            row += 2
        
        # 자연어 설명 섹션 추가
        natural_language_explanation = data.get("natural_language_explanation", "")
        if natural_language_explanation and natural_language_explanation.strip():
            ws[f'A{row}'] = "방책 추천 설명"
            ws[f'A{row}'].font = Font(bold=True, size=12)
            row += 1
            # 자연어 설명을 여러 줄로 분리하여 추가
            explanation_lines = natural_language_explanation.split('\n')
            for line in explanation_lines:
                if line.strip():
                    # 마크다운 헤더 제거
                    clean_line = line.strip().lstrip('#').strip()
                    if clean_line:
                        ws[f'A{row}'] = clean_line
                        row += 1
            row += 1
        
        # 추천 방책 요약 테이블
        ws[f'A{row}'] = "추천 방책 요약"
        ws[f'A{row}'].font = Font(bold=True, size=14)
        row += 1
        
        # 헤더
        headers = ["순위", "방책명", "적합도 점수"]
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col_idx)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = border
        row += 1
        
        # 데이터
        recommendations = data.get("recommendations", [])
        for idx, rec in enumerate(recommendations, 1):
            ws.cell(row=row, column=1).value = idx
            ws.cell(row=row, column=1).alignment = center_align
            ws.cell(row=row, column=1).border = border
            
            ws.cell(row=row, column=2).value = rec.get('coa_name', f'방책 {idx}')
            ws.cell(row=row, column=2).border = border
            
            ws.cell(row=row, column=3).value = rec.get('score', 0)
            ws.cell(row=row, column=3).alignment = center_align
            ws.cell(row=row, column=3).border = border
            row += 1
        
        # 상세 정보
        if include_details and recommendations:
            row += 1
            ws[f'A{row}'] = "추천 방책 상세"
            ws[f'A{row}'].font = Font(bold=True, size=14)
            row += 1
            
            for idx, rec in enumerate(recommendations, 1):
                ws[f'A{row}'] = f"{idx}. {rec.get('coa_name', f'방책 {idx}')}"
                ws[f'A{row}'].font = Font(bold=True)
                row += 1
                
                ws[f'A{row}'] = "적합도 점수:"
                ws[f'B{row}'] = rec.get('score', 0)
                row += 1
                
                if rec.get('description'):
                    ws[f'A{row}'] = "설명:"
                    ws[f'B{row}'] = rec.get('description', '')
                    row += 1
                
                # 점수 Breakdown
                score_breakdown = rec.get('score_breakdown', {})
                if score_breakdown:
                    row += 1
                    ws[f'A{row}'] = "점수 Breakdown"
                    ws[f'A{row}'].font = Font(bold=True)
                    row += 1
                    
                    # Breakdown 헤더
                    ws.cell(row=row, column=1).value = "요소"
                    ws.cell(row=row, column=1).fill = header_fill
                    ws.cell(row=row, column=1).font = header_font
                    ws.cell(row=row, column=1).border = border
                    ws.cell(row=row, column=2).value = "점수"
                    ws.cell(row=row, column=2).fill = header_fill
                    ws.cell(row=row, column=2).font = header_font
                    ws.cell(row=row, column=2).border = border
                    row += 1
                    
                    # Breakdown 데이터
                    for key, value in score_breakdown.items():
                        # METT-C 점수는 별도 섹션에서 처리
                        if key == 'mett_c':
                            continue
                        ws.cell(row=row, column=1).value = key
                        ws.cell(row=row, column=1).border = border
                        ws.cell(row=row, column=2).value = value
                        ws.cell(row=row, column=2).border = border
                        row += 1
                
                # METT-C 점수 (있는 경우)
                mett_c_scores = score_breakdown.get("mett_c") or rec.get("mett_c")
                if mett_c_scores:
                    row += 1
                    ws[f'A{row}'] = "METT-C 종합 평가"
                    ws[f'A{row}'].font = Font(bold=True, color="FFFFFF")
                    ws[f'A{row}'].fill = PatternFill(start_color="00b894", end_color="00b894", fill_type="solid")
                    row += 1
                    
                    # METT-C 헤더
                    ws.cell(row=row, column=1).value = "요소"
                    ws.cell(row=row, column=1).fill = header_fill
                    ws.cell(row=row, column=1).font = header_font
                    ws.cell(row=row, column=1).border = border
                    ws.cell(row=row, column=2).value = "점수"
                    ws.cell(row=row, column=2).fill = header_fill
                    ws.cell(row=row, column=2).font = header_font
                    ws.cell(row=row, column=2).border = border
                    ws.cell(row=row, column=3).value = "해석"
                    ws.cell(row=row, column=3).fill = header_fill
                    ws.cell(row=row, column=3).font = header_font
                    ws.cell(row=row, column=3).border = border
                    row += 1
                    
                    # METT-C 데이터
                    mett_c_elements = {
                        "mission": ("🎯 임무", 0.8),
                        "enemy": ("⚠️ 적군", 0.6),
                        "terrain": ("🌍 지형", 0.6),
                        "troops": ("👥 부대", 0.6),
                        "civilian": ("🏘️ 민간인", 0.3),
                        "time": ("⏰ 시간", 0.5)
                    }
                    
                    for key, (label, threshold) in mett_c_elements.items():
                        score = mett_c_scores.get(key, 0)
                        if score >= threshold:
                            interpretation = "양호"
                        elif score >= threshold * 0.5:
                            interpretation = "보통"
                        else:
                            interpretation = "부족"
                        
                        # 민간인/시간 특별 표시
                        if key == "civilian" and score < 0.3:
                            interpretation = "⚠️ 민간인 보호 낮음"
                        elif key == "time" and score == 0.0:
                            interpretation = "❌ 시간 제약 위반"
                        elif key == "time" and score < 0.5:
                            interpretation = "⚠️ 시간 제약 주의"
                        
                        ws.cell(row=row, column=1).value = label
                        ws.cell(row=row, column=1).border = border
                        ws.cell(row=row, column=2).value = score
                        ws.cell(row=row, column=2).border = border
                        ws.cell(row=row, column=3).value = interpretation
                        ws.cell(row=row, column=3).border = border
                        row += 1
                    
                    # METT-C 종합 점수
                    mett_c_total = mett_c_scores.get("total", 0)
                    ws[f'A{row}'] = "METT-C 종합 점수:"
                    ws[f'A{row}'].font = Font(bold=True)
                    ws[f'B{row}'] = mett_c_total
                    row += 1
                
                row += 1
        
        return row
    
    def _build_situation_excel_content(self, ws, data: Dict, start_row: int, header_fill, header_font, border, center_align) -> int:
        """상황 분석 보고서 Excel 내용"""
        row = start_row
        
        situation_info = data.get("situation_info", {})
        
        # 헤더
        headers = ["항목", "내용"]
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col_idx)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = center_align
            cell.border = border
        row += 1
        
        # 데이터
        for key, value in situation_info.items():
            ws.cell(row=row, column=1).value = key
            ws.cell(row=row, column=1).border = border
            ws.cell(row=row, column=2).value = str(value) if value is not None else "N/A"
            ws.cell(row=row, column=2).border = border
            row += 1
        
        return row
    
    def _build_rationale_excel_content(self, ws, data: Dict, start_row: int, header_fill, header_font, border, center_align) -> int:
        """의사결정 근거 보고서 Excel 내용"""
        row = start_row
        
        # 추론 단계
        reasoning_steps = data.get("reasoning_steps", [])
        if reasoning_steps:
            ws[f'A{row}'] = "추론 단계"
            ws[f'A{row}'].font = Font(bold=True, size=14)
            row += 1
            
            for step in reasoning_steps:
                ws[f'A{row}'] = step.get('step', '')
                ws[f'B{row}'] = step.get('description', '')
                row += 1
            row += 1
        
        # 점수 계산
        score_calculation = data.get("score_calculation", {})
        if score_calculation:
            ws[f'A{row}'] = "점수 계산"
            ws[f'A{row}'].font = Font(bold=True, size=14)
            row += 1
            
            for key, value in score_calculation.items():
                ws[f'A{row}'] = key
                ws[f'B{row}'] = value
                row += 1
        
        return row
    
    def _build_execution_excel_content(self, ws, data: Dict, start_row: int, header_fill, header_font, border, center_align) -> int:
        """실행 계획서 Excel 내용"""
        from openpyxl.styles import Font
        
        row = start_row
        
        coa = data.get("coa", {})
        execution_steps = data.get("execution_steps", [])
        
        # 방책 정보
        ws[f'A{row}'] = "방책명:"
        ws[f'B{row}'] = coa.get('coa_name', 'N/A')
        row += 1
        
        # 실행 단계
        if execution_steps:
            ws[f'A{row}'] = "실행 단계"
            ws[f'A{row}'].font = Font(bold=True, size=14)
            row += 1
            
            headers = ["순서", "단계", "설명"]
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col_idx)
                cell.value = header
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = border
            row += 1
            
            for idx, step in enumerate(execution_steps, 1):
                ws.cell(row=row, column=1).value = idx
                ws.cell(row=row, column=1).alignment = center_align
                ws.cell(row=row, column=1).border = border
                ws.cell(row=row, column=2).value = step.get('step', '') if isinstance(step, dict) else str(step)
                ws.cell(row=row, column=2).border = border
                ws.cell(row=row, column=3).value = step.get('description', '') if isinstance(step, dict) else ''
                ws.cell(row=row, column=3).border = border
                row += 1
        
        return row

