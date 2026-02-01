# api/routers/report.py
"""
보고서 생성 API - PDF, DOCX, TXT 형식 지원
"""
from fastapi import APIRouter, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import logging
import os
from datetime import datetime
from io import BytesIO

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/report", tags=["report"])

# 보고서 저장 디렉토리
REPORT_DIR = "reports"


class ReportRequest(BaseModel):
    """보고서 생성 요청 스키마"""
    agent_name: Optional[str] = "COA Recommendation Agent"
    summary: Optional[str] = None
    citations: Optional[List[str]] = []
    threat_info: Optional[Dict[str, Any]] = {}
    coa_recommendations: Optional[List[Dict[str, Any]]] = []
    format: str = "txt"  # "pdf", "docx", or "txt"


def _ensure_report_dir():
    """보고서 디렉토리 생성"""
    if not os.path.exists(REPORT_DIR):
        os.makedirs(REPORT_DIR)


def _generate_report_content(request: ReportRequest) -> Dict[str, Any]:
    """보고서 내용 구성 (공통)"""
    threat = request.threat_info or {}
    coas = request.coa_recommendations or []
    
    # 첫 번째 COA를 선정 방책으로 사용
    selected_coa = coas[0] if coas else {}
    
    content = {
        "title": "작전 분석 결과 보고서",
        "timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "threat": {
            "id": threat.get('threat_id', 'Unknown'),
            "type": threat.get('threat_type_original') or threat.get('threat_type', 'Unknown'),
            "description": threat.get('raw_report_text', '정보 없음'),
            "level": threat.get('threat_level', 'Unknown'),
            "location": threat.get('location', 'Unknown'),
        },
        "selected_coa": {
            "name": selected_coa.get('coa_name', 'Unknown'),
            "score": selected_coa.get('total_score', 0),
            "description": selected_coa.get('description', '정보 없음'),
            "reasoning_trace": selected_coa.get('reasoning_trace', []),
            "execution_plan": selected_coa.get('execution_plan', {}),
        },
        "summary": request.summary or "방책 추천 요약 정보 없음",
        "citations": request.citations or [],
        "all_coas": coas,
    }
    return content


def _generate_txt_report(content: Dict[str, Any], file_path: str):
    """TXT 보고서 생성"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(f"{'='*50}\n")
        f.write(f"         {content['title']}\n")
        f.write(f"{'='*50}\n")
        f.write(f"생성일시: {content['timestamp']}\n\n")
        
        # 1. 위협 상황
        f.write(f"1. 위협 상황 상세\n")
        f.write(f"   - 위협 ID: {content['threat']['id']}\n")
        f.write(f"   - 위협 유형: {content['threat']['type']}\n")
        f.write(f"   - 위협 수준: {content['threat']['level']}\n")
        f.write(f"   - 발생 위치: {content['threat']['location']}\n")
        f.write(f"   - 상황 설명: {content['threat']['description']}\n\n")
        
        # 2. 선정 방책
        f.write(f"2. 선정 방책 (Recommended COA)\n")
        f.write(f"   - 방책명: {content['selected_coa']['name']}\n")
        f.write(f"   - 종합 점수: {content['selected_coa']['score']}\n")
        f.write(f"   - 설명: {content['selected_coa']['description']}\n\n")
        
        # 3. 추론 근거
        f.write(f"3. 추론 근거\n")
        traces = content['selected_coa']['reasoning_trace']
        if traces:
            for i, trace in enumerate(traces):
                f.write(f"   Step {i+1}: {trace}\n")
        else:
            f.write(f"   (추론 근거 정보 없음)\n")
        f.write("\n")
        
        # 4. 실행 계획
        f.write(f"4. 실행 계획\n")
        plan = content['selected_coa']['execution_plan']
        phases = plan.get('phases', [])
        if phases:
            for i, phase in enumerate(phases):
                f.write(f"   Phase {i+1}: {phase.get('name', 'Unknown')}\n")
                f.write(f"      설명: {phase.get('description', '')}\n")
                for task in phase.get('tasks', []):
                    f.write(f"      - {task}\n")
        else:
            f.write(f"   (실행 계획 정보 없음)\n")
        f.write("\n")
        
        # 5. 요약
        f.write(f"5. 요약\n")
        f.write(f"   {content['summary']}\n\n")
        
        # 6. 참고 자료
        if content['citations']:
            f.write(f"6. 참고 자료\n")
            for i, citation in enumerate(content['citations']):
                f.write(f"   [{i+1}] {citation}\n")
        
        f.write(f"\n{'='*50}\n")
        f.write(f"         보고서 끝\n")
        f.write(f"{'='*50}\n")


def _generate_pdf_report(content: Dict[str, Any], file_path: str):
    """PDF 보고서 생성 (reportlab 사용)"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # 한글 폰트 등록 시도
        font_name = "Helvetica"
        try:
            # Windows 시스템 폰트 사용
            malgun_path = "C:/Windows/Fonts/malgun.ttf"
            if os.path.exists(malgun_path):
                pdfmetrics.registerFont(TTFont('Malgun', malgun_path))
                font_name = "Malgun"
        except Exception as e:
            logger.warning(f"한글 폰트 로드 실패, 기본 폰트 사용: {e}")
        
        doc = SimpleDocTemplate(file_path, pagesize=A4)
        styles = getSampleStyleSheet()
        
        # 커스텀 스타일 (한글 지원)
        title_style = ParagraphStyle(
            'TitleKR',
            parent=styles['Title'],
            fontName=font_name,
            fontSize=18,
            spaceAfter=20,
        )
        heading_style = ParagraphStyle(
            'HeadingKR',
            parent=styles['Heading2'],
            fontName=font_name,
            fontSize=14,
            spaceAfter=10,
        )
        body_style = ParagraphStyle(
            'BodyKR',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=10,
            spaceAfter=6,
        )
        
        story = []
        
        # 제목
        story.append(Paragraph(content['title'], title_style))
        story.append(Paragraph(f"생성일시: {content['timestamp']}", body_style))
        story.append(Spacer(1, 20))
        
        # 1. 위협 상황
        story.append(Paragraph("1. 위협 상황 상세", heading_style))
        threat_data = [
            ["항목", "내용"],
            ["위협 ID", content['threat']['id']],
            ["위협 유형", content['threat']['type']],
            ["위협 수준", str(content['threat']['level'])],
            ["발생 위치", content['threat']['location']],
        ]
        threat_table = Table(threat_data, colWidths=[100, 350])
        threat_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ]))
        story.append(threat_table)
        story.append(Paragraph(f"상황 설명: {content['threat']['description'][:500]}", body_style))
        story.append(Spacer(1, 15))
        
        # 2. 선정 방책
        story.append(Paragraph("2. 선정 방책 (Recommended COA)", heading_style))
        coa = content['selected_coa']
        story.append(Paragraph(f"<b>방책명:</b> {coa['name']}", body_style))
        story.append(Paragraph(f"<b>종합 점수:</b> {coa['score']}", body_style))
        story.append(Paragraph(f"<b>설명:</b> {coa['description'][:500]}", body_style))
        story.append(Spacer(1, 15))
        
        # 3. 요약
        story.append(Paragraph("3. 요약", heading_style))
        story.append(Paragraph(content['summary'][:1000], body_style))
        story.append(Spacer(1, 15))
        
        # 4. 참고 자료
        if content['citations']:
            story.append(Paragraph("4. 참고 자료", heading_style))
            for i, citation in enumerate(content['citations'][:10]):
                story.append(Paragraph(f"[{i+1}] {citation}", body_style))
        
        doc.build(story)
        
    except ImportError as e:
        logger.error(f"reportlab 라이브러리 오류: {e}")
        raise HTTPException(status_code=500, detail="PDF 생성 라이브러리(reportlab) 오류")
    except Exception as e:
        logger.error(f"PDF 생성 오류: {e}")
        raise HTTPException(status_code=500, detail=f"PDF 생성 실패: {str(e)}")


def _generate_docx_report(content: Dict[str, Any], file_path: str):
    """DOCX 보고서 생성 (python-docx 사용)"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 제목
        title = doc.add_heading(content['title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"생성일시: {content['timestamp']}")
        doc.add_paragraph()
        
        # 1. 위협 상황
        doc.add_heading("1. 위협 상황 상세", level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        threat_items = [
            ("위협 ID", content['threat']['id']),
            ("위협 유형", content['threat']['type']),
            ("위협 수준", str(content['threat']['level'])),
            ("발생 위치", content['threat']['location']),
            ("상황 설명", content['threat']['description'][:300]),
        ]
        for i, (label, value) in enumerate(threat_items):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
        
        doc.add_paragraph()
        
        # 2. 선정 방책
        doc.add_heading("2. 선정 방책 (Recommended COA)", level=1)
        coa = content['selected_coa']
        doc.add_paragraph(f"방책명: {coa['name']}")
        doc.add_paragraph(f"종합 점수: {coa['score']}")
        doc.add_paragraph(f"설명: {coa['description']}")
        doc.add_paragraph()
        
        # 3. 추론 근거
        doc.add_heading("3. 추론 근거", level=1)
        traces = coa['reasoning_trace']
        if traces:
            for i, trace in enumerate(traces):
                doc.add_paragraph(f"Step {i+1}: {trace}", style='List Number')
        else:
            doc.add_paragraph("(추론 근거 정보 없음)")
        doc.add_paragraph()
        
        # 4. 실행 계획
        doc.add_heading("4. 실행 계획", level=1)
        plan = coa['execution_plan']
        phases = plan.get('phases', [])
        if phases:
            for i, phase in enumerate(phases):
                doc.add_paragraph(f"Phase {i+1}: {phase.get('name', 'Unknown')}", style='Heading 3')
                doc.add_paragraph(f"설명: {phase.get('description', '')}")
                for task in phase.get('tasks', []):
                    doc.add_paragraph(f"• {task}", style='List Bullet')
        else:
            doc.add_paragraph("(실행 계획 정보 없음)")
        doc.add_paragraph()
        
        # 5. 요약
        doc.add_heading("5. 요약", level=1)
        doc.add_paragraph(content['summary'])
        doc.add_paragraph()
        
        # 6. 참고 자료
        if content['citations']:
            doc.add_heading("6. 참고 자료", level=1)
            for i, citation in enumerate(content['citations']):
                doc.add_paragraph(f"[{i+1}] {citation}")
        
        doc.save(file_path)
        
    except ImportError as e:
        logger.error(f"python-docx 라이브러리 오류: {e}")
        raise HTTPException(status_code=500, detail="DOCX 생성 라이브러리(python-docx) 오류")
    except Exception as e:
        logger.error(f"DOCX 생성 오류: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX 생성 실패: {str(e)}")


@router.post("/generate")
async def generate_report(request: ReportRequest):
    """
    분석 결과 보고서를 생성합니다.
    
    지원 형식:
    - txt: 텍스트 파일
    - pdf: PDF 파일 (reportlab 필요)
    - docx: Word 문서 (python-docx 필요)
    """
    try:
        _ensure_report_dir()
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"COA_Report_{timestamp}.{request.format}"
        file_path = os.path.join(REPORT_DIR, filename)
        
        # 보고서 내용 구성
        content = _generate_report_content(request)
        
        # 형식별 보고서 생성
        if request.format == "pdf":
            _generate_pdf_report(content, file_path)
            media_type = 'application/pdf'
        elif request.format == "docx":
            _generate_docx_report(content, file_path)
            media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:  # txt (기본)
            _generate_txt_report(content, file_path)
            media_type = 'text/plain; charset=utf-8'
        
        logger.info(f"Report generated: {file_path} (format: {request.format})")
        
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type=media_type
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating report: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"보고서 생성 실패: {str(e)}")
